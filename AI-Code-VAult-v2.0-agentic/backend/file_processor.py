# ============================================================================
# FILE PROCESSOR - File Upload & Text Extraction
# ============================================================================

import os
import re
from typing import List, Dict, Any
from pathlib import Path

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from various file types.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.csv':
            return extract_text_csv(file_path)
        elif ext == '.txt':
            return extract_text_plain(file_path)
        elif ext == '.pdf':
            return extract_text_pdf(file_path)
        elif ext == '.docx':
            return extract_text_docx(file_path)
        elif ext in ['.py', '.js', '.java', '.cpp', '.c', '.go', '.rb', '.php']:
            return extract_text_code(file_path)
        else:
            return extract_text_plain(file_path)
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def extract_text_csv(file_path: str) -> str:
    """Extract text from CSV using pandas. Each row becomes its own pre-formed chunk string."""
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        text = ""
        for _, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            # Format each row as a distinct block so chunk_text can split it or ingest_file_to_vault uses it directly
            text += row_text + "\n__CSV_ROW_BOUNDARY__\n"
        return text
    except ImportError:
        print("pandas not installed. Install with: pip install pandas")
        return ""
    except Exception as e:
        print(f"Error extracting CSV: {str(e)}")
        return ""

def extract_text_plain(file_path: str) -> str:
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading plain text: {str(e)}")
        return ""

def extract_text_code(file_path: str) -> str:
    """Extract text from code file."""
    return extract_text_plain(file_path)

def extract_text_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except ImportError:
        print("PyPDF2 not installed. Install with: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"Error extracting PDF: {str(e)}")
        return ""

def extract_text_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Error extracting DOCX: {str(e)}")
        return ""

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """
    Split text into overlapping chunks for processing.
    
    Args:
        text: Text to chunk
        chunk_size: Size of each chunk in characters
        overlap: Overlap between chunks
        
    Returns:
        List of text chunks
    """
    chunks = []
    
    if len(text) <= chunk_size:
        return [text]
    
    # Split by sentences first
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += sentence + " "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # Add overlap from previous chunk
            if len(chunks) > 0:
                current_chunk = chunks[-1][-overlap:] + " " + sentence + " "
            else:
                current_chunk = sentence + " "
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def validate_file(file_path: str, max_size_mb: int = 50) -> Dict[str, Any]:
    """
    Validate uploaded file.
    
    Args:
        file_path: Path to the file
        max_size_mb: Maximum file size in MB
        
    Returns:
        Validation result
    """
    if not os.path.exists(file_path):
        return {
            'valid': False,
            'error': 'File does not exist'
        }
    
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
    
    if file_size_mb > max_size_mb:
        return {
            'valid': False,
            'error': f'File size {file_size_mb:.2f}MB exceeds maximum {max_size_mb}MB'
        }
    
    # Check file extension
    supported_extensions = {'.txt', '.pdf', '.docx', '.py', '.js', '.java', '.cpp', '.c', '.go', '.rb', '.php', '.csv'}
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext not in supported_extensions:
        return {
            'valid': False,
            'error': f'File type {ext} not supported'
        }
    
    return {
        'valid': True,
        'size_mb': file_size_mb,
        'extension': ext
    }

def get_file_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        File metadata
    """
    stat = os.stat(file_path)
    
    return {
        'filename': os.path.basename(file_path),
        'size': stat.st_size,
        'created_at': stat.st_ctime,
        'modified_at': stat.st_mtime,
        'extension': os.path.splitext(file_path)[1].lower()
    }

def ingest_file_to_vault(file_path: str, filename: str, user_id: int, session, progress_callback=None):
    """
    Unified Ingestion Function implementing Data Vault 2.0.
    
    1. Creates Document row immediately.
    2. Creates HubDocument and LinkUserDocument.
    3. Chunks content (CSV row-by-row, or 1000/100 overlapping chars).
    4. Iterates chunks to create SatDocumentContent.
    5. Generates embedding and creates SatDocumentEmbedding.
    6. Commits every 10 chunks to prevent cloud timeouts.
    """
    from db_connector import Document, HubDocument, LinkUserDocument, SatDocumentContent, SatDocumentEmbedding, HubUser, User, SatUserProfile, load_data_vault, compute_hash_key, compute_hash_diff
    from embeddings import get_embedding as generate_embedding
    from datetime import datetime
    import time
    
    start_time = time.time()
    file_ext = os.path.splitext(filename)[1].lower()
    file_size = os.path.getsize(file_path)
    
    # 1. Immediate Document creation
    new_doc = Document(
        user_id=user_id,
        filename=filename,
        file_type=file_ext,
        size=file_size,
        upload_date=datetime.now(),
        chunk_count=0,
        status='processing'
    )
    session.add(new_doc)
    session.commit()
    
    try:
        # 2. Extract and chunk
        chunks = []
        if file_ext == '.csv':
            import pandas as pd
            df = pd.read_csv(file_path)
            # Store column names in new_doc metadata
            new_doc.column_names = [str(c) for c in df.columns]
            session.commit()
            for _, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                chunks.append(row_text)
        else:
            raw_text = extract_text_from_file(file_path)
            if raw_text:
                # Handle our internal CSV fallback format just in case
                if '__CSV_ROW_BOUNDARY__' in raw_text:
                    chunks = [c.strip() for c in raw_text.split('__CSV_ROW_BOUNDARY__') if c.strip()]
                else:
                    chunks = chunk_text(raw_text, chunk_size=1000, overlap=100)
                    
        total_chunks = len(chunks)
        
        if total_chunks == 0:
            new_doc.status = 'failed (no text)'
            session.commit()
            return {"chunks": 0, "embeddings_success": 0, "embeddings_failed": 0, "time": time.time() - start_time}

        if progress_callback:
            progress_callback(0, total_chunks)

        # 3. Create Hub and Link
        doc_hash = compute_hash_key(filename, user_id)
        if not session.query(HubDocument).filter_by(hash_key=doc_hash).first():
            session.add(HubDocument(hash_key=doc_hash, record_source='vault_app', filename=filename))
            session.flush()
            
        user = session.query(User).filter_by(id=user_id).first()
        if user:
            user_hash = compute_hash_key(user.email)
            if not session.query(HubUser).filter_by(hash_key=user_hash).first():
                session.add(HubUser(hash_key=user_hash, record_source='vault_app', email=user.email))
                session.flush()
                
            link_hash = compute_hash_key(user_hash, doc_hash)
            if not session.query(LinkUserDocument).filter_by(hash_key=link_hash).first():
                session.add(LinkUserDocument(hash_key=link_hash, record_source='vault_app', hub_user_hash=user_hash, hub_doc_hash=doc_hash))
                session.flush()

            # Ensure SatUserProfile is loaded / updated
            sat_user_data = {
                'role': user.role or 'User',
                'preferences': {'theme': 'System'}
            }
            load_data_vault(
                session=session,
                hub_model=HubUser,
                satellite_model=SatUserProfile,
                hub_data={'email': user.email},
                sat_data=sat_user_data,
                hub_hash_key=user_hash,
                record_source='vault_app',
                hub_fk_column='hub_user_hash'
            )

        embeddings_success = 0
        embeddings_failed = 0
        
        # 4. Iterate every single chunk
        for i, chunk_str in enumerate(chunks):
            chunk_hash_diff = compute_hash_diff(raw_text=chunk_str, chunk_index=i)
            
            # Create SatDocumentContent
            sat_content = SatDocumentContent(
                hub_doc_hash=doc_hash,
                load_date=datetime.now(),
                hash_diff=chunk_hash_diff,
                raw_text=chunk_str,
                chunk_index=i,
                chunk_size=len(chunk_str)
            )
            session.add(sat_content)
            
            # Generate embedding with try/except
            emb_vector = None
            try:
                emb_vector = generate_embedding(chunk_str)
                embeddings_success += 1
            except Exception as e:
                print(f"Embedding generation failed for chunk {i}: {e}")
                embeddings_failed += 1
                
            # Create SatDocumentEmbedding
            sat_emb = SatDocumentEmbedding(
                hub_doc_hash=doc_hash,
                load_date=datetime.now(),
                hash_diff=chunk_hash_diff,
                embedding_vector=emb_vector,
                model_name='sentence-transformers'
            )
            session.add(sat_emb)
            
            # Commit every 10 chunks to ensure partial ingestion on cloud timeouts
            if (i + 1) % 10 == 0:
                session.commit()
                if progress_callback:
                    progress_callback(i + 1, total_chunks)
                    
        # Final commit and document update
        new_doc.chunk_count = total_chunks
        new_doc.status = 'complete'
        session.commit()
        
        if progress_callback:
            progress_callback(total_chunks, total_chunks)
            
        return {
            "chunks": total_chunks,
            "embeddings_success": embeddings_success,
            "embeddings_failed": embeddings_failed,
            "time": time.time() - start_time
        }
        
    except Exception as e:
        session.rollback()
        new_doc.status = 'failed'
        session.commit()
        print(f"Ingestion failed: {e}")
        raise e
